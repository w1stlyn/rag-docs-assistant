from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".txt"}


@dataclass
class Document:
    text: str
    source: str
    metadata: dict = field(default_factory=dict)


def load_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in {".md", ".txt"}:
        return _load_plain(path)
    raise ValueError(f"Неподдерживаемый формат: {suffix}")


def load_directory(path: Path) -> list[Document]:
    docs: list[Document] = []
    for file in sorted(path.rglob("*")):
        if file.is_file() and file.suffix.lower() in SUPPORTED_SUFFIXES:
            docs.extend(load_file(file))
    return docs


def _load_pdf(path: Path) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(Document(text=text, source=path.name, metadata={"path": str(path), "page": i}))
    return docs


def _load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDoc

    doc = DocxDoc(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text:
        return []
    return [Document(text=text, source=path.name, metadata={"path": str(path), "page": 1})]


def _load_plain(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8")
    if not text.strip():
        return []
    return [Document(text=text, source=path.name, metadata={"path": str(path), "page": 1})]
